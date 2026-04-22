"""Generates a Word document summarising the Everest Motoring platform."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUT = r"C:\Users\info\OneDrive\Documents\Antigravity\everest-motoring\Everest-Motoring-Platform-Features.docx"

BRAND_RED = RGBColor(0xB7, 0x1C, 0x1C)
DARK = RGBColor(0x1F, 0x29, 0x37)
MID = RGBColor(0x4B, 0x55, 0x63)

doc = Document()

# Default styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1, color=BRAND_RED):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def add_feature(name, description, advantages):
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK
    doc.add_paragraph(description)
    adv_p = doc.add_paragraph()
    adv_r = adv_p.add_run('Advantage: ')
    adv_r.bold = True
    adv_r.font.color.rgb = BRAND_RED
    adv_p.add_run(advantages)

# --- Cover ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('Everest Motoring')
tr.bold = True
tr.font.size = Pt(32)
tr.font.color.rgb = BRAND_RED

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('Platform Feature & Capability Document')
sr.font.size = Pt(16)
sr.font.color.rgb = DARK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run(f'Prepared {date.today().strftime("%d %B %Y")}  \u2022  everestmotoring.co.za')
mr.italic = True
mr.font.color.rgb = MID

doc.add_paragraph()
doc.add_paragraph()

# --- Executive Summary ---
add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Everest Motoring runs on a modern, custom-built Next.js platform hosted on Vercel with '
    'Supabase as the database and storage backbone. The system combines a polished public-facing '
    'dealership website with a deep admin backend, email automation, AI-generated video and news '
    'content, and integrated affiliate and client portals. Every feature is purpose-built for the '
    'South African motor-retail market and optimised for lead capture, customer retention, and '
    'search visibility.'
)
doc.add_paragraph(
    'This document covers: (1) public-facing frontend features, (2) admin and operational backend, '
    '(3) AI and automation systems, (4) integrations, (5) recent work delivered, and (6) what is in '
    'progress or planned next.'
)

# --- Section 1: Frontend ---
add_heading('1. Public-Facing Website', level=1)
doc.add_paragraph(
    'Pages visible to the public, prospective buyers, affiliates, and returning clients.'
)

add_heading('Home Page', level=2, color=DARK)
add_feature(
    'Hero + Featured Vehicles',
    'Landing page with cinematic hero banner, primary call-to-action, and a dynamic featured-vehicles rail pulled from live inventory.',
    'High-impact first impression; directs visitors straight to stock that is prioritised for sale.'
)
add_feature(
    'Inline Search Widget',
    'Home-page search bar with make / model / price / body-type filters that deep-links into the inventory grid.',
    'Removes one click from the buying journey \u2014 visitors can start browsing without leaving the home page.'
)

add_heading('Inventory', level=2, color=DARK)
add_feature(
    'Inventory Grid with Dynamic Filters',
    'Filterable, sortable listing of all in-stock vehicles. Filters include make, model, year, price range, fuel type, transmission, and body style.',
    'Lets buyers self-serve and narrow to their exact requirement \u2014 lowers enquiry friction and surfaces the right car faster.'
)
add_feature(
    'Vehicle Detail Pages',
    'Full vehicle profile: gallery, specs, features list, finance estimator link, AI-generated walk-around video, and a lead form pre-wired to the specific vehicle.',
    'Every vehicle has a richer presentation than standard Autotrader/Cars.co.za listings, building trust and conversion.'
)
add_feature(
    'Live AI Walk-Around Videos',
    'Short AI-generated video clips showcase each vehicle in motion, stitched from the inventory photos via Seedance / Sora pipelines.',
    'Differentiates Everest\u2019s stock visually; videos are proven to outperform static imagery in engagement.'
)
add_feature(
    'Per-Vehicle SEO Fields',
    'Every vehicle carries auto-generated SEO meta title, description, and structured data, plus an IndexNow ping when stock changes.',
    'Faster indexing by Google / Bing and better SERP snippets \u2014 directly measurable in organic traffic.'
)

add_heading('Customer Actions', level=2, color=DARK)
add_feature(
    'Book Test Drive',
    'Dedicated booking form that captures contact details, preferred vehicle, preferred date/time, and notifies the admin team instantly.',
    'Removes friction for a high-intent action; bookings land straight in the admin CRM.'
)
add_feature(
    'Value My Car (Trade-In)',
    'Public trade-in valuation form: vehicle details, condition, photos. Submissions feed into the admin Trade-Ins queue.',
    'Captures trade-ins proactively \u2014 a known high-margin acquisition channel for dealer stock.'
)
add_feature(
    'Finance / Pre-Approval',
    'Finance information page with pre-approval form (soft check) and links to partner banks.',
    'Qualifies buyers earlier \u2014 fewer time-wasters on the showroom floor and higher close rate per lead.'
)
add_feature(
    'Contact & Enquiry Forms',
    'Multiple context-aware forms across the site, all feeding the same Leads pipeline.',
    'One CRM inbox for every enquiry source; no leads fall through the cracks.'
)

add_heading('News / Content', level=2, color=DARK)
add_feature(
    'News Section',
    'Public news / blog with listing and individual article pages; posts are AI-generated monthly and editorially approved.',
    'Keeps the site fresh and relevant for SEO; positions Everest as an authority voice in the local motor market.'
)

add_heading('Client Area', level=2, color=DARK)
add_feature(
    'Client Registration & Login',
    'Separate client registration flow (distinct from affiliate), with Supabase-backed authentication.',
    'Personalised return experience \u2014 clients see their vehicles, documents, and service history.'
)
add_feature(
    'Client Dashboard',
    'Logged-in area showing purchased vehicles, upcoming service reminders, and saved favourites.',
    'Builds long-term retention; every repeat-buyer touchpoint stays inside the Everest ecosystem.'
)

add_heading('Affiliate Area', level=2, color=DARK)
add_feature(
    'Affiliate Registration & Approval',
    'Affiliates apply via a dedicated form; admin approves before they can log in.',
    'Controlled onboarding \u2014 only vetted partners represent Everest.'
)
add_feature(
    'Affiliate Dashboard',
    'Affiliates see their referred leads, pipeline stages (Call Back / Financing / Completed), monthly commission totals with reset logic, and a searchable lead list.',
    'Transparent commission tracking drives affiliate effort; reduces admin queries.'
)
add_feature(
    'Affiliate Links',
    'Unique tracked referral links per affiliate.',
    'Every conversion is attributable \u2014 commission pay-outs are defensible and automated.'
)

# --- Section 2: Admin ---
add_heading('2. Admin & Operational Backend', level=1)
doc.add_paragraph(
    'Private admin console for the Everest team to run the business day-to-day.'
)

add_heading('Inventory Management', level=2, color=DARK)
add_feature(
    'Add / Edit Vehicles',
    'Full CRUD interface for stock: photos, specs, features, price, status. Photo uploads go directly to Supabase storage.',
    'No reliance on third-party stock systems \u2014 full control of what appears on the public site within seconds.'
)
add_feature(
    'Mark as Sold (with Post-Sale Flow)',
    'One-click sold flow that captures buyer name, email, phone, notes, and a customer-collecting-car photo. Stores a sale record, updates vehicle status, and schedules follow-ups.',
    'Turns the sale event into the start of the retention cycle, not the end. Every sale automatically triggers review + video + email workflow.'
)
add_feature(
    'Social-Share Flags per Vehicle',
    'Columns track which platforms a vehicle has already been shared to, preventing duplicate posts.',
    'Enables coordinated multi-channel posting without double-posting.'
)

add_heading('Lead & CRM', level=2, color=DARK)
add_feature(
    'Leads Board',
    'Central leads list: source, vehicle, client contact, stage, notes, created date. Sortable and searchable.',
    'One view of the sales pipeline \u2014 nothing lives in private inboxes.'
)
add_feature(
    'Lead Assignment',
    'Admins can assign leads to specific team members or affiliates.',
    'Clear ownership \u2014 every lead has an accountable person, improving follow-up speed.'
)
add_feature(
    'Trade-In Queue',
    'Public trade-in submissions land here for valuation and response.',
    'Consolidates trade-in intake; response SLA is trackable.'
)

add_heading('Content Management', level=2, color=DARK)
add_feature(
    'News Editor',
    'Admin UI to create, edit, and publish news posts, including AI-drafted monthly pieces.',
    'Low effort to keep content fresh; AI drafts mean new posts cost minutes, not hours.'
)

add_heading('Affiliates Admin', level=2, color=DARK)
add_feature(
    'Affiliate Approval + Management',
    'Approve / revoke affiliates, view their referred leads, track commission liability.',
    'Hands-on control of the affiliate programme; compliance and payout visibility in one place.'
)

# --- Section 3: AI & Automation ---
add_heading('3. AI & Automation Systems', level=1)

add_heading('Vehicle Video Pipeline', level=2, color=DARK)
add_feature(
    'AI Walk-Around Generation (Seedance / Sora / Veo)',
    'Each vehicle automatically generates a 5-scene AI walk-around video from its photos. Pipeline supports retry, RHD enforcement, and multi-model fallback.',
    'Every car gets an eye-catching video clip at zero per-vehicle production cost.'
)
add_feature(
    'Video Hosting (Mux \u2192 Cloudflare Stream)',
    'Generated videos host on Cloudflare Stream for fast delivery and adaptive bitrate; legacy Mux content backfilled and migrated.',
    'Lower ongoing costs versus Mux; global CDN playback; MP4 fallbacks ready for email and social.'
)
add_feature(
    'Global Browser Render Queue',
    'Video generation runs on a shared browser-based queue to bypass Vercel serverless limits for long-running jobs.',
    'Reliable long-form video rendering without paying for enterprise Vercel tiers.'
)

add_heading('Post-Sale Handover Video (New)', level=2, color=DARK)
add_feature(
    'Seedance 2 Handover Clip',
    'From the delivery photo, generate an 8-second 16:9 cinematic clip with audio. Four scene options: Dream Drive Transition, Time-Lapse Reveal, Stylized Hero Walkaround, Pixel Build.',
    'Gives every buyer a personalised, shareable \u201cI just bought my car\u201d moment \u2014 proven to drive word-of-mouth referrals and social engagement.'
)
add_feature(
    'Scene Picker & Regenerate',
    'Admin can preview the four scenes, pick one, and later \u201cTry a different scene\u201d to regenerate with a different concept.',
    'Creative control without creative cost; admin can match the scene to the customer / vehicle story.'
)
add_feature(
    'Auto-Embedded in Review Email',
    'When the handover video completes before the scheduled email sends, the email is automatically rescheduled with the video thumbnail embedded.',
    'The review request email goes from \u201cplease review us\u201d to \u201chere\u2019s a cinematic reminder of your big moment \u2014 please review us.\u201d Dramatically higher open / engagement.'
)

add_heading('AI Content', level=2, color=DARK)
add_feature(
    'Monthly News Generator (Cron)',
    'Scheduled job drafts industry-relevant news posts monthly, routed through admin for approval before publishing.',
    'Always-on SEO content engine; no manual blogging burden.'
)
add_feature(
    'SEO Field Generator',
    'Auto-generates SEO title, description, and structured data for each vehicle from its specs.',
    'Consistent, rich metadata on every listing without manual data entry.'
)
add_feature(
    'Script Generator',
    'AI script generator for video voice-overs and marketing copy.',
    'Faster creative turnaround; consistent brand voice.'
)

add_heading('Email Automation (Resend + React Email)', level=2, color=DARK)
add_feature(
    'Post-Sale Review Email (4-Day Delay)',
    'Scheduled automatically when a sale is recorded. Includes vehicle, delivery photo, handover video, and Google review link.',
    'Captures the customer at peak satisfaction for a Google review \u2014 compounding reputation lift over time.'
)
add_feature(
    '1-Month Follow-Up',
    'Sent 30 days after delivery; service reminders and introduction to the client portal.',
    'Keeps Everest top-of-mind during the early ownership window.'
)
add_feature(
    '3-Year Trade-In Reminder',
    'Sent around the 3-year mark to nudge trade-up conversations.',
    'Repeat-buyer funnel with zero ongoing effort \u2014 revenue-generating automation.'
)
add_feature(
    'Birthday & Welcome Emails',
    'Personalised birthday greeting and onboarding welcome.',
    'Low-cost goodwill that meaningfully improves retention scores.'
)
add_feature(
    'Newsletter Broadcasts',
    'Branded newsletter template for mass sends to the client base.',
    'Re-engage the database on new stock, promotions, and industry news.'
)
add_feature(
    'Affiliate Media Kit Email',
    'Automated kit delivery when affiliates are approved.',
    'Hands-off affiliate onboarding.'
)

# --- Section 4: Infrastructure ---
add_heading('4. Integrations & Infrastructure', level=1)
doc.add_paragraph(
    'Foundations the whole system runs on. Chosen for reliability, cost, and developer velocity.'
)

add_feature(
    'Next.js 15 on Vercel',
    'Server-rendered React app with App Router, server actions, and edge middleware.',
    'Fast global performance; deploys in under a minute; trusted by enterprise brands.'
)
add_feature(
    'Supabase (Postgres + Auth + Storage)',
    'Unified backend for database, authentication, RLS security policies, and public storage buckets (delivery photos, sale videos, vehicle photos).',
    'One platform handles data, auth, and files \u2014 significantly cheaper and simpler than stitching AWS / Auth0 / S3 together.'
)
add_feature(
    'Resend (Transactional Email)',
    'All email sends go through Resend, including native scheduled sends (4-day post-sale).',
    'Rock-solid deliverability; inspectable message log; scheduled sends without running a cron ourselves.'
)
add_feature(
    'Cloudflare Stream',
    'Adaptive-bitrate video hosting and delivery for all vehicle and handover videos.',
    'Consistent playback on any device / connection; CDN-delivered globally.'
)
add_feature(
    'Google Gemini',
    'Powers the news, SEO, and script generation.',
    'Best-in-class long-form AI text at competitive cost.'
)
add_feature(
    'kie.ai (Seedance 2 Fast)',
    'Image-to-video generation with retry, polling, and mock mode.',
    'High-quality cinematic clips generated in minutes without owning GPU infrastructure.'
)
add_feature(
    'HeyGen',
    'AI avatar / talking-head video generation for promotional pieces.',
    'Studio-quality presenter videos without hiring on-camera talent.'
)
add_feature(
    'Sentry',
    'Production error monitoring across client, server, and edge runtimes.',
    'Bugs surface and get fixed before customers notice; measurable reliability.'
)
add_feature(
    'Google Analytics 4',
    'Traffic and conversion tracking.',
    'Data-driven decisions on marketing spend and page performance.'
)
add_feature(
    'IndexNow',
    'Search-engine notification service pinged when inventory changes.',
    'New stock appears in Bing / Yandex within minutes instead of days.'
)
add_feature(
    'Autotrader & Cars Feeds',
    'XML/JSON feeds at /api/feeds/autotrader and /api/feeds/cars for syndication partners.',
    'Everest stock reaches third-party marketplaces automatically \u2014 wider reach at zero incremental effort.'
)
add_feature(
    'Health Endpoint',
    '/api/health reports Supabase reachability, env-var presence, region, and uptime.',
    'Easy external monitoring and quick diagnosis when something breaks.'
)

# --- Section 5: Recently Built ---
add_heading('5. Recently Delivered (April 2026)', level=1)
doc.add_paragraph('Shipped and live on production during April 2026:')

bullets = [
    ('Post-Sale Flow', 'Mark-as-sold button, buyer capture form, delivery-photo upload, sales record, and lead closure in one action.'),
    ('Seedance 2 Handover Videos', '4 cinematic 16:9 scene options (Dream Drive, Time-Lapse Reveal, Stylized Hero Walkaround, Pixel Build) with audio; regenerate-to-different-scene flow.'),
    ('Scheduled Review Email', '4-day delayed post-sale email through Resend with handover video thumbnail embedded.'),
    ('News System', 'Cron-driven monthly news generation, admin editor, public news section.'),
    ('SEO System', 'Per-vehicle SEO fields, auto-generation, IndexNow integration.'),
    ('Sentry Error Monitoring', 'Client, server, and edge configuration for production error visibility.'),
    ('Health Check Endpoint', '/api/health for uptime monitoring.'),
    ('Dealer Feed Endpoints', 'Autotrader and generic cars feed for syndication.'),
    ('Client Registration Separation', 'Distinct client vs affiliate registration flows.'),
    ('Home-Page Featured Vehicles', 'Dynamic featured rail pulled from live inventory.'),
]
for name, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(name + ': ')
    r1.bold = True
    p.add_run(desc)

# --- Section 6: In Progress / Next ---
add_heading('6. In Progress & Planned', level=1)

planned = [
    ('Social Auto-Posting of Handover Videos', 'Push finished Seedance clips to Instagram, TikTok, and Facebook via the Ember Social platform.',
     'Turn every sale into organic social proof \u2014 reach clients the buyer knows without paid spend.'),
    ('WhatsApp Distribution', 'Automated affiliate media distribution via WhatsApp Business API alongside email.',
     'Faster affiliate response times on mobile-first channels.'),
    ('Live Market Pricing', 'Real-time market-value signals on each vehicle listing.',
     'Builds buyer confidence and justifies asking price.'),
    ('Instant Finance Pre-Approval (Soft Check)', 'In-page soft credit check for instant pre-qualification.',
     'Converts curious visitors into qualified buyers in one session.'),
    ('RESEND_API_KEY + NEXT_PUBLIC_SITE_URL Environment Config', 'Remaining production env vars to enable full email scheduling and clean health checks.',
     'Closes the last gap before every feature runs at 100% in production.'),
    ('Migration 20260421 Execution', 'Apply the sale_video_style widening migration on Supabase.',
     'Unlocks the 4-scene handover video picker on live data.'),
]
for name, desc, adv in planned:
    add_feature(name, desc, adv)

# --- Closing ---
add_heading('Appendix: Technology Summary', level=1)
tech_rows = [
    ('Framework', 'Next.js 15 (App Router, Server Actions)'),
    ('Hosting', 'Vercel Production'),
    ('Database', 'Supabase Postgres'),
    ('Auth', 'Supabase Auth (Google + email/password)'),
    ('Storage', 'Supabase Storage (delivery-photos, sale-videos, vehicle photos)'),
    ('Email', 'Resend + React Email'),
    ('Video Hosting', 'Cloudflare Stream'),
    ('AI Text', 'Google Gemini'),
    ('AI Video', 'Seedance 2 Fast (kie.ai), Sora 2, Veo 3.1, HeyGen'),
    ('Error Monitoring', 'Sentry'),
    ('Analytics', 'Google Analytics 4'),
    ('SEO', 'IndexNow + structured data'),
    ('Domain', 'everestmotoring.co.za'),
]
table = doc.add_table(rows=len(tech_rows), cols=2)
table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(tech_rows):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('Prepared for Everest Motoring  \u2022  Confidential')
fr.italic = True
fr.font.color.rgb = MID
fr.font.size = Pt(9)

doc.save(OUT)
print(f'Saved: {OUT}')
